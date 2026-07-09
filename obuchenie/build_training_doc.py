from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = Path("training_output")
IMG_OUT = OUT / "annotated"
OUT.mkdir(exist_ok=True)
IMG_OUT.mkdir(exist_ok=True)


BLUE = "2E74B5"
DARK = "0B2545"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
GREEN = "EAF7ED"
GOLD = "FFF3CD"
RED = "FDECEC"
BORDER = "DADCE0"


def font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def annotate(src, dst, crop=(0, 58, 960, 590), markers=(), footer=None):
    im = Image.open(src).convert("RGB").crop(crop)
    pad = 24
    header_h = 64
    footer_h = 42 if footer else 0
    canvas = Image.new("RGB", (im.width + pad * 2, im.height + pad * 2 + header_h + footer_h), "white")
    d = ImageDraw.Draw(canvas)
    d.rounded_rectangle([0, 0, canvas.width - 1, canvas.height - 1], radius=18, fill="white", outline=(210, 218, 226), width=2)
    d.rectangle([0, 0, canvas.width - 1, header_h], fill=(11, 37, 69))
    d.text((pad, 18), "Скриншот из Selsup: зона действия выделена номером", fill="white", font=font(22, True))
    canvas.paste(im, (pad, pad + header_h))
    for n, x, y, label, color in markers:
        x2 = x - crop[0] + pad
        y2 = y - crop[1] + pad + header_h
        fill = color
        d.ellipse([x2 - 18, y2 - 18, x2 + 18, y2 + 18], fill=fill, outline="white", width=3)
        w = d.textlength(str(n), font=font(22, True))
        d.text((x2 - w / 2, y2 - 13), str(n), fill="white", font=font(22, True))
        if label:
            label_x = min(max(x2 + 24, 16), canvas.width - 300)
            label_y = min(max(y2 - 18, header_h + 8), canvas.height - 76)
            d.rounded_rectangle([label_x, label_y, label_x + 280, label_y + 36], radius=8, fill=(255, 255, 255), outline=(46, 116, 181), width=2)
            d.text((label_x + 10, label_y + 8), label, fill=(11, 37, 69), font=font(17, True))
    if footer:
        y = canvas.height - footer_h
        d.rectangle([1, y, canvas.width - 2, canvas.height - 2], fill=(244, 246, 249))
        d.text((pad, y + 10), footer, fill=(72, 82, 94), font=font(18))
    canvas.save(dst)
    return dst


figures = {
    "product_toolbar": annotate(
        "video_frames/frame_009.png",
        IMG_OUT / "01_product_toolbar.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 127, 141, "Открыта карточка товара", (46, 116, 181)),
            (2, 374, 140, "Иконка Честного знака", (245, 158, 11)),
            (3, 158, 575, "Сохранить", (20, 160, 185)),
        ],
        footer="Используйте верхнюю панель карточки, чтобы перейти к связям маркетплейсов и маркировке.",
    ),
    "error_settings": annotate(
        "video_frames/frame_010.png",
        IMG_OUT / "02_error_settings.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 705, 68, "Ошибка GTIN / запроса", (210, 50, 45)),
            (2, 78, 253, "Настройки", (46, 116, 181)),
        ],
        footer="Если Selsup сообщает об ошибке запроса GTIN, проверьте интеграцию Честного знака.",
    ),
    "integrations": annotate(
        "video_frames/frame_011.png",
        IMG_OUT / "03_integrations.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 123, 112, "Настройки -> Интеграции", (46, 116, 181)),
            (2, 574, 183, "Честный знак", (245, 158, 11)),
            (3, 578, 253, "Настроить", (20, 160, 185)),
        ],
        footer="В списке интеграций найдите блок Честного знака и откройте настройку.",
    ),
    "token": annotate(
        "video_frames/frame_012.png",
        IMG_OUT / "04_token.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 130, 182, "Организация", (46, 116, 181)),
            (2, 171, 371, "Получить идентификатор", (20, 160, 185)),
            (3, 166, 486, "Получить токен", (245, 158, 11)),
            (4, 299, 486, "Проверить СУЗ", (36, 151, 81)),
        ],
        footer="После получения токена сохраните настройку и убедитесь, что СУЗ подключен.",
    ),
    "success": annotate(
        "video_frames/frame_014.png",
        IMG_OUT / "05_success.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 780, 63, "Успешное обновление", (36, 151, 81)),
            (2, 388, 377, "Данные карточки", (46, 116, 181)),
            (3, 151, 574, "Сохранить", (20, 160, 185)),
        ],
        footer="После возврата в карточку дождитесь зелёного уведомления и сохраните изменения.",
    ),
    "status_table": annotate(
        "video_frames/frame_016.png",
        IMG_OUT / "06_status_table.png",
        crop=(0, 64, 960, 588),
        markers=[
            (1, 180, 243, "Статус по размеру", (46, 116, 181)),
            (2, 482, 243, "GTIN", (245, 158, 11)),
            (3, 591, 243, "Штрихкод", (36, 151, 81)),
            (4, 151, 574, "Сохранить", (20, 160, 185)),
        ],
        footer="Финальная проверка: по каждому размеру есть статус, GTIN и штрихкод.",
    ),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    style_run(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        style_run(p.add_run(text))
    return p


def add_callout(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "E0E6ED")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, 11, True, DARK)
    p.add_run("\n")
    r2 = p.add_run(body)
    style_run(r2, 10, False, DARK)
    return table


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    style_run(r, 9, False, "555555")


def add_step(doc, num, title, why, action, result, image_key):
    add_heading(doc, f"Шаг {num}. {title}", 2)
    add_callout(doc, "Зачем это нужно", why, LIGHT_BLUE)
    add_body(doc, "Действие: " + action, "Действие:")
    add_body(doc, "Результат: " + result, "Результат:")
    add_figure(doc, figures[image_key], f"Иллюстрация к шагу {num}: {title.lower()}.")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.text = "База знаний склада | Selsup и Честный знак"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_run(header.runs[0], 9, False, "555555")
footer = section.footer.paragraphs[0]
footer.text = "Учебный пример для адаптации сотрудника"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.runs[0], 9, False, "555555")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Учебный модуль\n")
style_run(r, 28, True, DARK)
r = title.add_run("Честный знак в Selsup: завершение модерации и обновление карточки товара")
style_run(r, 18, True, BLUE)

subtitle = doc.add_paragraph()
r = subtitle.add_run("Формат: текстово-графическая инструкция по видео. Аудитория: новый сотрудник склада / оператор карточек. Версия: пилотный пример.")
style_run(r, 11, False, "555555")

add_callout(
    doc,
    "Что сотрудник должен уметь после урока",
    "Понять, где в Selsup находится интеграция Честного знака, как обновить токен/СУЗ, вернуться в карточку товара и проверить, что по размерам подтянулись GTIN и штрихкоды.",
    GREEN,
)

add_heading(doc, "Карта процесса", 1)
table = doc.add_table(rows=1, cols=4)
table.autofit = False
set_table_width(table, [0.8, 2.1, 2.7, 1.9])
headers = ["Этап", "Где работаем", "Что делаем", "Готово, если"]
for i, h in enumerate(headers):
    c = table.rows[0].cells[i]
    set_cell_shading(c, LIGHT_BLUE)
    set_cell_border(c)
    r = c.paragraphs[0].add_run(h)
    style_run(r, 10, True, DARK)
rows = [
    ("1", "Карточка товара", "Проверяем статус и признаки проблемы с GTIN/маркировкой.", "Понятно, нужна ли настройка ЧЗ."),
    ("2", "Настройки -> Интеграции", "Открываем блок Честного знака.", "Карточка интеграции найдена."),
    ("3", "Раздел Честный знак", "Выбираем организацию, сертификат, получаем токен и проверяем СУЗ.", "Появились зелёные статусы."),
    ("4", "Карточка товара", "Возвращаемся к товару, сохраняем и дожидаемся обновления.", "Есть зелёное уведомление."),
    ("5", "Таблица размеров", "Проверяем статус, GTIN и штрихкод по каждому размеру.", "Нет пустых GTIN/штрихкодов."),
]
for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        set_cell_border(cells[i])
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = cells[i].paragraphs[0].add_run(val)
        style_run(r, 9.5, i == 0, DARK)

add_heading(doc, "Перед началом", 1)
add_body(doc, "Доступ: сотрудник вошёл в Selsup под учётной записью с правами на товары, настройки и интеграции.", "Доступ:")
add_body(doc, "Карточка товара: товар уже создан, у него есть варианты/размеры и данные для маркетплейсов.", "Карточка товара:")
add_body(doc, "ЭЦП/сертификат: сертификат организации доступен на рабочем месте, где выполняется настройка.", "ЭЦП/сертификат:")
add_body(doc, "Важно: не удаляйте GTIN и штрихкоды вручную, если не уверены в причине ошибки. Сначала проверьте интеграцию.", "Важно:")

doc.add_page_break()

add_step(
    doc,
    1,
    "Открыть карточку товара и оценить проблему",
    "Перед настройкой нужно понять, что проблема относится именно к маркировке/GTIN, а не к обычному заполнению карточки.",
    "Откройте нужный товар в разделе “Товары”. Посмотрите верхнюю панель карточки и блоки с параметрами: категория, артикул, бренд, страна производства, название для этикетки.",
    "Карточка открыта, сотрудник видит товар и понимает, какие данные будут обновляться.",
    "product_toolbar",
)

add_step(
    doc,
    2,
    "Перейти в настройки при ошибке GTIN",
    "Красное уведомление о запросе GTIN означает, что Selsup не смог корректно получить или обновить данные маркировки.",
    "Если появилось уведомление об ошибке, откройте левое меню и перейдите в “Настройки”. Не закрывайте карточку без сохранения важных изменений.",
    "Сотрудник перешёл к настройкам и готов проверить интеграцию Честного знака.",
    "error_settings",
)

add_step(
    doc,
    3,
    "Открыть интеграцию Честного знака",
    "Интеграция отвечает за обмен с Честным знаком: токен, СУЗ и получение кодов/идентификаторов.",
    "В разделе “Интеграции” найдите карточку “Честный знак”. Если на ней статус “частично настроено”, нажмите “Настроить”.",
    "Открыт экран настройки Честного знака для выбранной организации.",
    "integrations",
)

doc.add_page_break()

add_step(
    doc,
    4,
    "Получить токен и проверить СУЗ",
    "Без актуального токена Selsup не сможет стабильно обращаться к Честному знаку и обновлять данные по карточкам.",
    "Выберите организацию и сертификат ЭЦП. Нажмите “Получить идентификатор соединения”, затем “Получить токен”. Укажите категорию товара и нажмите “Проверить СУЗ”. После успешной проверки нажмите “Сохранить”.",
    "На экране видны зелёные статусы: токен получен, СУЗ успешно подключён, информация сохранена.",
    "token",
)

add_step(
    doc,
    5,
    "Вернуться в карточку и сохранить обновление",
    "После настройки интеграции карточка должна подтянуть новые данные и зафиксировать их в Selsup.",
    "Вернитесь в карточку товара. Проверьте основные поля и нажмите “Сохранить”. Дождитесь зелёного уведомления об успешном обновлении карточки.",
    "Selsup показывает, что карточка/информация успешно обновлена.",
    "success",
)

add_step(
    doc,
    6,
    "Проверить размеры, GTIN и штрихкоды",
    "Финальная ошибка часто находится не в общей карточке, а в конкретном размере: пустой GTIN, неправильный размер или не подтянутый штрихкод.",
    "Прокрутите карточку до таблицы размеров. По каждой строке проверьте статус, российский размер, размер производителя, GTIN и штрихкод. Если всё заполнено, сохраните карточку.",
    "По всем размерам есть статусы и заполненные идентификаторы. Карточку можно использовать дальше в процессах поставки.",
    "status_table",
)

doc.add_page_break()
add_heading(doc, "Контроль качества", 1)
checks = [
    "В карточке товара нет красного уведомления об ошибке GTIN.",
    "В настройках Честного знака выбран правильный сертификат организации.",
    "Токен получен, а СУЗ проверен успешно.",
    "После возврата в карточку появилось зелёное уведомление об обновлении.",
    "В таблице размеров по каждой строке заполнены GTIN и штрихкод.",
    "Карточка сохранена после финальной проверки.",
]
for item in checks:
    p = doc.add_paragraph(style="List Bullet")
    style_run(p.add_run(item), 11, False, DARK)

add_heading(doc, "Частые ошибки", 1)
mistakes = doc.add_table(rows=1, cols=3)
mistakes.autofit = False
set_table_width(mistakes, [2.0, 3.0, 2.5])
for i, h in enumerate(["Ситуация", "Что это значит", "Что сделать"]):
    c = mistakes.rows[0].cells[i]
    set_cell_shading(c, LIGHT_BLUE)
    set_cell_border(c)
    style_run(c.paragraphs[0].add_run(h), 10, True, DARK)
for row in [
    ("Красное уведомление GTIN", "Selsup не получил данные от Честного знака или не смог повторить запрос.", "Проверить интеграцию, токен и СУЗ."),
    ("Статус “частично настроено”", "Интеграция есть, но не все обязательные данные подтверждены.", "Открыть настройку и пройти получение токена."),
    ("Пустой GTIN в одном размере", "Обновление прошло не по всем строкам таблицы размеров.", "Проверить конкретный размер, затем сохранить карточку."),
    ("Нет зелёного уведомления", "Сохранение или обновление не завершилось.", "Подождать, обновить страницу и повторить сохранение."),
]:
    cells = mistakes.add_row().cells
    for i, val in enumerate(row):
        set_cell_border(cells[i])
        r = cells[i].paragraphs[0].add_run(val)
        style_run(r, 9.5, i == 0, DARK)

add_heading(doc, "Как использовать этот формат дальше", 1)
add_callout(
    doc,
    "Шаблон для следующих видео",
    "Для каждого процесса склада можно собирать такой же модуль: цель урока, карта процесса, 5-8 шагов со скриншотами, контроль качества и частые ошибки. Видео остаётся источником, а документ становится быстрым рабочим регламентом для новичка.",
    PALE,
)

output = OUT / "Учебный модуль - Честный знак в Selsup.docx"
doc.save(output)
print(output.resolve())
