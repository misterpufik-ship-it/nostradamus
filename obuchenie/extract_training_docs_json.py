import json
from pathlib import Path

from docx import Document


ROOT = Path.cwd()
targets = []

for file in (ROOT / "Videos").glob("*.docx"):
    if file.name.startswith("~$"):
        continue
    if "Поставка на WB" in file.name or file.name == "Поставка.docx":
        targets.append(file)

for file in ROOT.parent.glob("*.docx"):
    if file.name.startswith("~$"):
        continue
    if "дополнения к регламенту" in file.name:
        targets.append(file)

items = []
for path in targets:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(values)
        if rows:
            tables.append(rows)
    items.append({"path": str(path), "paragraphs": paragraphs, "tables": tables})

out = ROOT / "training_output" / "wb_docs_extracted.json"
out.parent.mkdir(exist_ok=True)
out.write_text(json.dumps(items, ensure_ascii=True, indent=2), encoding="ascii")
print(out)
