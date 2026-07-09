from pathlib import Path

from docx import Document


ROOT = Path.cwd()
targets = []

for file in (ROOT / "Videos").glob("*.docx"):
    if file.name.startswith("~$"):
        continue
    if "Поставка на WB" in file.name or "Поставка" == file.stem:
        targets.append(file)

desktop = ROOT.parent
for file in desktop.glob("*.docx"):
    if file.name.startswith("~$"):
        continue
    if "дополнения к регламенту" in file.name:
        targets.append(file)

out = ROOT / "training_output" / "wb_docs_extracted.txt"
out.parent.mkdir(exist_ok=True)

lines = []
for path in targets:
    lines.append(f"### FILE: {path}")
    doc = Document(path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(doc.tables, 1):
        lines.append(f"--- TABLE {table_index}")
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(values):
                lines.append(" || ".join(values))
    lines.append("")

out.write_text("\n".join(lines), encoding="utf-8")
print(out)
